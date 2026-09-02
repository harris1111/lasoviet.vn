from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from pathlib import Path


OUT = Path('/workspace/scratch/205957f1e170/La-So-Viet-Brand-Experience-Guideline-v1.0.docx')

# Brand palette
PAPER_50 = 'FFFDF7'
PAPER_100 = 'F7F1E5'
PAPER_200 = 'EEE5D6'
PAPER_300 = 'D7CDBD'
CONTROL_BORDER = '958A7C'
INK_900 = '14263D'
INK_800 = '263445'
INK_600 = '5E6873'
CINNABAR = 'A63D2F'
CINNABAR_DARK = '81271F'
CINNABAR_LIGHT = 'F3DDD4'
LINK_BLUE = '174F7A'
GREEN = '2F6F57'
AMBER = '8A5A12'
ERROR = '8F2737'
WHITE = 'FFFFFF'

UI_FONT = 'DejaVu Sans'
DISPLAY_FONT = 'DejaVu Serif'
BRAND_UI_FONT = 'Be Vietnam Pro Variable'
BRAND_DISPLAY_FONT = 'Source Serif 4 Variable'


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(width_dxa))
    tc_w.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tbl_pr.append(layout)
    layout.set(qn('w:type'), 'fixed')
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(widths_dxa)))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.find(qn('w:tblInd'))
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), str(indent_dxa))
    tbl_ind.set(qn('w:type'), 'dxa')

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[min(idx, len(widths_dxa) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement('w:tblHeader')
    header.set(qn('w:val'), 'true')
    tr_pr.append(header)


def set_font(run, font=UI_FONT, size=None, bold=None, italic=None, color=None):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), font)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), font)
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text, url, color=LINK_BLUE):
    part = paragraph.part
    rel_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), rel_id)
    run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')
    r_fonts = OxmlElement('w:rFonts')
    r_fonts.set(qn('w:ascii'), UI_FONT)
    r_fonts.set(qn('w:hAnsi'), UI_FONT)
    r_pr.append(r_fonts)
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    r_pr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    r_pr.append(u)
    run.append(r_pr)
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.extend([fld_char1, instr, fld_char2])
    set_font(run, UI_FONT, 8.5, color=INK_600)


def add_num_defs(doc):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn('w:abstractNumId'))) for x in numbering.findall(qn('w:abstractNum'))]
    existing_num = [int(x.get(qn('w:numId'))) for x in numbering.findall(qn('w:num'))]
    next_abs = max(existing_abs or [0]) + 1
    next_num = max(existing_num or [0]) + 1

    def make_abstract(abs_id, fmt, text, font=None):
        abstract = OxmlElement('w:abstractNum')
        abstract.set(qn('w:abstractNumId'), str(abs_id))
        multi = OxmlElement('w:multiLevelType')
        multi.set(qn('w:val'), 'singleLevel')
        abstract.append(multi)
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), '0')
        start = OxmlElement('w:start'); start.set(qn('w:val'), '1'); lvl.append(start)
        num_fmt = OxmlElement('w:numFmt'); num_fmt.set(qn('w:val'), fmt); lvl.append(num_fmt)
        lvl_text = OxmlElement('w:lvlText'); lvl_text.set(qn('w:val'), text); lvl.append(lvl_text)
        suff = OxmlElement('w:suff'); suff.set(qn('w:val'), 'tab'); lvl.append(suff)
        p_pr = OxmlElement('w:pPr')
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab'); tab.set(qn('w:val'), 'num'); tab.set(qn('w:pos'), '540'); tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement('w:ind'); ind.set(qn('w:left'), '540'); ind.set(qn('w:hanging'), '280'); p_pr.append(ind)
        spacing = OxmlElement('w:spacing'); spacing.set(qn('w:after'), '80'); spacing.set(qn('w:line'), '300'); spacing.set(qn('w:lineRule'), 'auto'); p_pr.append(spacing)
        lvl.append(p_pr)
        if font:
            r_pr = OxmlElement('w:rPr')
            fonts = OxmlElement('w:rFonts'); fonts.set(qn('w:ascii'), font); fonts.set(qn('w:hAnsi'), font); r_pr.append(fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

    make_abstract(next_abs, 'bullet', '•', UI_FONT)
    bullet_num = next_num
    num = OxmlElement('w:num'); num.set(qn('w:numId'), str(bullet_num))
    abs_ref = OxmlElement('w:abstractNumId'); abs_ref.set(qn('w:val'), str(next_abs)); num.append(abs_ref); numbering.append(num)

    make_abstract(next_abs + 1, 'decimal', '%1.')
    decimal_num = next_num + 1
    num2 = OxmlElement('w:num'); num2.set(qn('w:numId'), str(decimal_num))
    abs_ref2 = OxmlElement('w:abstractNumId'); abs_ref2.set(qn('w:val'), str(next_abs + 1)); num2.append(abs_ref2); numbering.append(num2)
    return bullet_num, decimal_num


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn('w:numPr'))
    if num_pr is None:
        num_pr = OxmlElement('w:numPr')
        p_pr.append(num_pr)
    ilvl = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), '0')
    numid = OxmlElement('w:numId'); numid.set(qn('w:val'), str(num_id))
    num_pr.extend([ilvl, numid])


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)
section.header_distance = Inches(0.38)
section.footer_distance = Inches(0.38)
section.different_first_page_header_footer = True

styles = doc.styles

normal = styles['Normal']
normal.font.name = UI_FONT
normal._element.rPr.rFonts.set(qn('w:ascii'), UI_FONT)
normal._element.rPr.rFonts.set(qn('w:hAnsi'), UI_FONT)
normal.font.size = Pt(10.4)
normal.font.color.rgb = RGBColor.from_string(INK_800)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.22

for style_name, font_name, size, color, before, after in [
    ('Title', DISPLAY_FONT, 30, INK_900, 0, 8),
    ('Subtitle', UI_FONT, 13, INK_600, 0, 8),
    ('Heading 1', DISPLAY_FONT, 19, INK_900, 16, 7),
    ('Heading 2', DISPLAY_FONT, 14.5, INK_900, 12, 5),
    ('Heading 3', UI_FONT, 11.5, CINNABAR, 9, 4),
]:
    st = styles[style_name]
    st.font.name = font_name
    st._element.rPr.rFonts.set(qn('w:ascii'), font_name)
    st._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = style_name != 'Subtitle'
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

# Major chapters always begin on a new page. Using the heading property instead
# of inserting a standalone page-break paragraph prevents trailing blank pages
# when the previous chapter ends close to the bottom margin.
styles['Heading 1'].paragraph_format.page_break_before = True

# Some office renderers apply a theme border to the built-in Title style.
# Remove it so the cover stays within the Paper / Ink / Cinnabar system.
title_ppr = styles['Title']._element.get_or_add_pPr()
title_border = title_ppr.find(qn('w:pBdr'))
if title_border is not None:
    title_ppr.remove(title_border)

if 'Lead' not in styles:
    lead = styles.add_style('Lead', WD_STYLE_TYPE.PARAGRAPH)
else:
    lead = styles['Lead']
lead.font.name = DISPLAY_FONT
lead._element.rPr.rFonts.set(qn('w:ascii'), DISPLAY_FONT)
lead._element.rPr.rFonts.set(qn('w:hAnsi'), DISPLAY_FONT)
lead.font.size = Pt(13.5)
lead.font.color.rgb = RGBColor.from_string(INK_900)
lead.paragraph_format.space_after = Pt(10)
lead.paragraph_format.line_spacing = 1.25

if 'Caption Small' not in styles:
    caption = styles.add_style('Caption Small', WD_STYLE_TYPE.PARAGRAPH)
else:
    caption = styles['Caption Small']
caption.font.name = UI_FONT
caption._element.rPr.rFonts.set(qn('w:ascii'), UI_FONT)
caption._element.rPr.rFonts.set(qn('w:hAnsi'), UI_FONT)
caption.font.size = Pt(8.5)
caption.font.color.rgb = RGBColor.from_string(INK_600)
caption.paragraph_format.space_after = Pt(4)

bullet_num, decimal_num = add_num_defs(doc)


def add_body(text='', bold_prefix=None, italic=False, color=None, align=None, keep=False):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, UI_FONT, 10.4, bold=True, color=color or INK_800)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2, UI_FONT, 10.4, italic=italic, color=color or INK_800)
    else:
        r = p.add_run(text)
        set_font(r, UI_FONT, 10.4, italic=italic, color=color or INK_800)
    if align is not None:
        p.alignment = align
    p.paragraph_format.keep_together = keep
    return p


def add_lead(text):
    p = doc.add_paragraph(style='Lead')
    p.add_run(text)
    return p


def add_h1(text):
    return doc.add_paragraph(text, style='Heading 1')


def add_h2(text):
    return doc.add_paragraph(text, style='Heading 2')


def add_h3(text):
    return doc.add_paragraph(text, style='Heading 3')


def add_bullets(items):
    for item in items:
        p = doc.add_paragraph()
        apply_num(p, bullet_num)
        if isinstance(item, tuple):
            label, rest = item
            r = p.add_run(label)
            set_font(r, UI_FONT, 10.3, bold=True, color=INK_900)
            r2 = p.add_run(rest)
            set_font(r2, UI_FONT, 10.3, color=INK_800)
        else:
            r = p.add_run(item)
            set_font(r, UI_FONT, 10.3, color=INK_800)


def add_numbered(items):
    for item in items:
        p = doc.add_paragraph()
        apply_num(p, decimal_num)
        r = p.add_run(item)
        set_font(r, UI_FONT, 10.3, color=INK_800)


def add_callout(label, text, fill=PAPER_200, accent=CINNABAR, text_color=INK_900):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.08)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); p_pr.append(shd)
    borders = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    for key, val in [('val', 'single'), ('sz', '18'), ('space', '8'), ('color', accent)]:
        left.set(qn(f'w:{key}'), val)
    borders.append(left); p_pr.append(borders)
    r1 = p.add_run(label.upper() + '  ')
    set_font(r1, UI_FONT, 9.2, bold=True, color=accent)
    r2 = p.add_run(text)
    set_font(r2, UI_FONT, 10.4, color=text_color)
    return p


def add_table(headers, rows, widths, header_fill=INK_900, header_color=WHITE, first_col_bold=False, font_size=9.1, row_fills=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_font(r, UI_FONT, font_size, bold=True, color=header_color)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        fill = row_fills[ridx] if row_fills and ridx < len(row_fills) else (PAPER_50 if ridx % 2 == 0 else WHITE)
        for i, value in enumerate(row):
            set_cell_shading(cells[i], fill)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            r = p.add_run(str(value))
            set_font(r, UI_FONT, font_size, bold=(first_col_bold and i == 0), color=INK_800)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def page_break():
    # Chapter pagination is controlled by Heading 1. Retained as a semantic
    # marker in the authoring script so content remains easy to reorganize.
    return None


# Running header/footer
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hr = hp.add_run('LÁ SỐ VIỆT  /  BRAND & EXPERIENCE GUIDELINE')
set_font(hr, UI_FONT, 8.4, bold=True, color=INK_600)
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fr = fp.add_run('v1.0  •  31.08.2026   |   ')
set_font(fr, UI_FONT, 8.3, color=INK_600)
add_page_field(fp)

# Cover — editorial cover pattern
for _ in range(3):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('LÁ SỐ VIỆT')
set_font(r, UI_FONT, 11, bold=True, color=CINNABAR)
p.paragraph_format.space_after = Pt(14)
p = doc.add_paragraph(style='Title')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Brand & Experience\nGuideline')
p.paragraph_format.space_after = Pt(10)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Nền tảng lập và luận giải lá số dành cho người Việt')
set_font(r, DISPLAY_FONT, 14.5, italic=True, color=INK_600)
p.paragraph_format.space_after = Pt(28)
add_callout('North Star', 'Thư viện tri thức Việt đương đại: tĩnh, sáng rõ, có căn cứ và trả quyền lựa chọn về cho người dùng.', fill=PAPER_100, accent=CINNABAR)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('VERSION 1.0  •  MVP / PHASE 1  •  31 AUGUST 2026')
set_font(r, UI_FONT, 9, bold=True, color=INK_600)
p.paragraph_format.space_before = Pt(30)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Dành cho Founder, Product, Design, Content, Growth và Engineering')
set_font(r, UI_FONT, 9.2, color=INK_600)

# Start body content in its own section. This preserves a clean editorial cover
# while keeping running headers/footers stable on every subsequent page.
content_section = doc.add_section(WD_SECTION.NEW_PAGE)
content_section.page_width = Inches(8.5)
content_section.page_height = Inches(11)
content_section.top_margin = Inches(0.85)
content_section.bottom_margin = Inches(0.8)
content_section.left_margin = Inches(0.9)
content_section.right_margin = Inches(0.9)
content_section.header_distance = Inches(0.38)
content_section.footer_distance = Inches(0.38)
content_section.different_first_page_header_footer = False
content_section.header.is_linked_to_previous = False
content_section.footer.is_linked_to_previous = False

body_header = content_section.header
body_hp = body_header.paragraphs[0]
body_hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
body_hr = body_hp.add_run('LÁ SỐ VIỆT  /  BRAND & EXPERIENCE GUIDELINE')
set_font(body_hr, UI_FONT, 8.4, bold=True, color=INK_600)
body_footer = content_section.footer
body_fp = body_footer.paragraphs[0]
body_fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
body_fr = body_fp.add_run('v1.0  •  31.08.2026   |   ')
set_font(body_fr, UI_FONT, 8.3, color=INK_600)
add_page_field(body_fp)

add_h1('Cách dùng guideline này')
add_lead('Đây là “hiến pháp thương hiệu” cho Lá Số Việt. Mọi quyết định về sản phẩm, nội dung, giao diện, growth và vận hành phải truy được về một nguyên tắc trong tài liệu này.')
add_table(
    ['Lớp quyết định', 'Ý nghĩa', 'Cách xử lý'],
    [
        ['LOCKED — Bất biến', 'Định vị, đạo đức, lời hứa, trust/safety và nguyên tắc không định mệnh hóa.', 'Chỉ đổi khi founder cập nhật decision log với lý do và bằng chứng.'],
        ['STANDARD — Chuẩn hệ thống', 'Màu, font, component, voice, layout và pattern đã đủ tốt để triển khai.', 'Tuân thủ mặc định; ngoại lệ phải có tên, lý do và owner.'],
        ['HYPOTHESIS — Cần test', 'Thông điệp challenger, hành vi “người Việt”, cách hiển thị mobile, giá và conversion.', 'Đo bằng user research/analytics; không biến thành chân lý thương hiệu.'],
    ],
    [1800, 3600, 3960], first_col_bold=True, font_size=8.8
)
add_callout('Quy tắc ưu tiên', 'Khi conversion xung đột với niềm tin, quyền riêng tư hoặc sự an toàn tâm lý, trust/safety thắng. Không dùng nỗi sợ để đổi lấy doanh thu ngắn hạn.', fill=CINNABAR_LIGHT, accent=CINNABAR_DARK)

add_h2('Mục lục')
for item in [
    '01. Brand foundation — nền tảng chiến lược',
    '02. Audience & behavioral truth — người dùng và hành vi',
    '03. Mood, personality & creative concept',
    '04. Tone of voice & messaging system',
    '05. Visual identity system',
    '06. Product experience principles',
    '07. Trust, privacy & safety by design',
    '08. Component and content recipes',
    '09. Validation roadmap & metrics',
    '10. Governance & decision checklist',
    'Appendix A. Copy library',
    'Appendix B. Evidence & sources',
]:
    p = doc.add_paragraph()
    r = p.add_run(item)
    set_font(r, UI_FONT, 10.2, bold=True if item.startswith(('01','02','03','04','05','06','07','08','09','10')) else False, color=INK_900)
    p.paragraph_format.space_after = Pt(3)

page_break()

add_h1('01. Brand foundation — nền tảng chiến lược')
add_h2('1.1 Brand role')
add_callout('Quyết định khóa', 'Lá Số Việt là người dẫn giải có căn cứ — không phải “thầy bói AI”, không phải công cụ tính toán lạnh lùng và không phải cỗ máy phán tương lai.')
add_body('Vai trò của thương hiệu là giúp người dùng hiểu cấu trúc lá số, đối chiếu với trải nghiệm thực tế và nhìn thấy điều họ có thể chủ động. Thương hiệu không thay người dùng ra quyết định.')

add_h2('1.2 Brand essence')
add_lead('Hiểu mình có căn cứ.')
add_body('Essence này cô đọng bốn giá trị: cá nhân hóa thật, diễn giải rõ, truy nguyên được và không tước quyền tự quyết. Nó phải hiện diện trong sản phẩm ngay cả khi logo, chiến dịch hoặc giao diện thay đổi.')

add_h2('1.3 Purpose, vision, mission')
add_table(
    ['Tầng', 'Tuyên bố'],
    [
        ['Purpose', 'Giúp người Việt biến sự phức tạp của lá số thành một góc nhìn rõ ràng để tự hiểu và chủ động hơn.'],
        ['Vision', 'Trở thành nền tảng tin cậy nhất tại Việt Nam để khám phá nhiều hệ quy chiếu Đông–Tây trên một hồ sơ sinh, với phương pháp riêng của từng hệ được tôn trọng.'],
        ['Mission — Phase 1', 'Lập lá số Tử Vi chính xác theo rule set công bố; tạo trải nghiệm miễn phí có giá trị; bán chiều sâu bằng báo cáo rõ căn cứ, an toàn và giao ngay online.'],
    ], [1800, 7560], first_col_bold=True
)

add_h2('1.4 Positioning statement')
add_callout('Định vị', 'Dành cho người Việt đang tìm sự rõ ràng về bản thân, quan hệ hoặc một giai đoạn nhiều bất định, Lá Số Việt là nền tảng lập và luận giải lá số bằng tiếng Việt rõ ràng, truy nguyên được căn cứ. Khác với trải nghiệm coi bói dựa vào uy quyền hoặc nỗi sợ, Lá Số Việt tách bạch tính toán, quy tắc diễn giải và AI, đồng thời trả quyền quyết định về cho người dùng.', fill=PAPER_100, accent=INK_900)

add_h2('1.5 Brand promise & reasons to believe')
add_bullets([
    ('Lời hứa chức năng: ', 'lập lá số theo dữ liệu đầu vào và phương pháp được công bố; báo cáo có cấu trúc và dễ đọc.'),
    ('Lời hứa cảm xúc: ', 'người dùng cảm thấy được tôn trọng, bớt mơ hồ và vẫn giữ quyền lựa chọn.'),
    ('Lời hứa đạo đức: ', 'không phán định, không dọa để upsell, không giả chuyên gia, không giấu AI.'),
    ('Bằng chứng sản phẩm: ', 'engine/version, evidence key, “Vì sao có nhận định này?”, giới hạn giờ sinh, report bất biến và dữ liệu riêng tư.'),
])

add_h2('1.6 Brand architecture')
add_bullets([
    ('Master brand: ', 'Lá Số Việt — luôn đủ dấu; domain canonical là lasoviet.vn.'),
    ('Category: ', 'nền tảng lập và luận giải lá số.'),
    ('Hiện tại: ', 'Tử Vi là paid MVP và lời hứa sản phẩm đang có.'),
    ('Tương lai: ', 'Bát Tự, Bản đồ sao, Kinh Dịch và tổng hợp Đông–Tây chỉ được truyền thông như sản phẩm hiện hữu khi engine/evidence đã sẵn sàng.'),
    ('AI: ', 'là cơ chế tổ chức và diễn giải; không nằm trong tên thương hiệu hoặc hero.'),
])

add_h2('1.7 What we are / What we are not')
add_table(
    ['Lá Số Việt là', 'Lá Số Việt không là'],
    [
        ['Một nền tảng tri thức và tự chiêm nghiệm', 'Một “thầy” có quyền phán quyết'],
        ['Cá nhân hóa dựa trên dữ liệu và evidence', 'Văn bản Barnum chung chung khoác áo cá nhân'],
        ['Tĩnh, rõ, ấm và hiện đại', 'Tím neon, vũ trụ, khói và sân khấu huyền bí'],
        ['Minh bạch về AI và giới hạn', 'Khoa học giả hoặc tuyên bố chính xác 99%'],
        ['Bán chiều sâu và cấu trúc', 'Bán sự giải thoát khỏi nỗi sợ'],
    ], [4680, 4680]
)

page_break()

add_h1('02. Audience & behavioral truth — người dùng và hành vi')
add_h2('2.1 Phân khúc theo tình huống, không theo stereotype')
add_body('Repo không chứng minh tuổi, giới, thu nhập hoặc vùng miền của người mua. Vì vậy guideline không dựng “chân dung người Việt” từ trực giác. Phân khúc đầu tiên dựa trên Job-to-be-Done và trạng thái bất định:')
add_table(
    ['Job / Moment', 'Pain chính', 'Giá trị cần giao'],
    [
        ['Hiểu bản thân & đường đời', 'Lá số khó đọc; mô tả trên mạng quá chung.', 'Bản đồ rõ cấu trúc, insight cá nhân, căn cứ và câu hỏi phản tư.'],
        ['Tình yêu & hôn nhân', 'Lo âu, muốn biết điều gì đang diễn ra hoặc có nên tiếp tục.', 'Ngôn ngữ không phán xét, pattern quan hệ, giới hạn và hành động lành mạnh.'],
        ['Công việc & tài lộc', 'Thiếu khung tham chiếu cho lựa chọn hoặc giai đoạn mới.', 'Xu hướng, điều kiện, rủi ro quan sát được; không đưa lời khuyên đầu tư.'],
        ['Vận trình năm', 'Muốn chuẩn bị trước nhưng dễ hiểu thành dự báo chắc chắn.', 'Cửa sổ thời gian, dấu hiệu, checklist chuẩn bị và mức độ tin cậy.'],
    ], [2200, 3100, 4060], font_size=8.7
)

add_h2('2.2 Những gì đã có bằng chứng')
add_bullets([
    ('Method-first acquisition: ', 'demand tập trung ở “tử vi”, “lá số tử vi”, “bát tự”, “bản đồ sao”, “kinh dịch”; landing/SEO phải gọi đúng tên phương pháp.'),
    ('Evidence-first conversion: ', 'người dùng cần thấy chart, sample, phương pháp và giá trị trước khi mua chiều sâu.'),
    ('Trust là biến số kinh doanh: ', 'nghiên cứu trên người mua online tại Việt Nam cho thấy website quality, service quality, information security và reference groups liên quan đáng kể đến trust; trust liên hệ mạnh với quyết định mua [S5].'),
    ('Bối cảnh chống lừa đảo là thật: ', 'Bộ Công an ghi nhận hơn 6.000 vụ việc lừa đảo trực tuyến năm 2024, thiệt hại hơn 12.000 tỷ đồng [S6].'),
    ('Thanh toán số đã phổ biến nhưng an toàn vẫn quan trọng: ', 'NAPAS ghi nhận tăng trưởng mạnh của VietQR trong 2025, đồng thời nhấn mạnh fraud prevention và bảo mật [S7].'),
    ('Niềm tin Tử Vi có hệ quả xã hội: ', 'nghiên cứu CESifo về matching hôn nhân tại Việt Nam cho thấy tính “cát” theo Tử Vi liên quan đến lựa chọn và hỗ trợ từ gia đình; đây là bằng chứng rằng sản phẩm cần trách nhiệm cao, không phải bằng chứng mọi người Việt đều tin [S4].'),
])

add_h2('2.3 Những gì chỉ là giả thuyết')
add_bullets([
    'Người dùng Việt ưu tiên authority cue, gia đình, thể diện hoặc cộng đồng ở mức nào trong trải nghiệm tử vi online.',
    'Người trẻ xem để giải trí còn người lớn tuổi xem để quyết định; nữ giới tin hoặc trả tiền nhiều hơn nam giới.',
    'Serif, giấy ngà hoặc màu son tự động làm tăng trust; giọng “thầy” phù hợp với mọi vùng miền.',
    'Người dùng muốn báo cáo dài hơn thay vì bản tóm tắt hành động; họ hiểu “có xu hướng” khác “sẽ xảy ra”.',
])
add_callout('Nguyên tắc research', 'Không tìm thấy khảo sát đại diện đáng tin về hành vi dùng tử vi online tại Việt Nam. Mọi nhận định văn hóa chưa có dữ liệu phải gắn nhãn “hypothesis” và được kiểm chứng bằng phỏng vấn, usability test hoặc traffic thật.', fill=AMBER+'22' if False else PAPER_200, accent=AMBER)

add_h2('2.4 Behavioral risks cần thiết kế chủ động')
add_table(
    ['Cơ chế', 'Rủi ro', 'Quy tắc UX'],
    [
        ['Barnum / personal validation', 'Câu chung chung được cảm nhận là “rất đúng với tôi”.', 'Mỗi insight gắn evidence; nêu điều kiện, giới hạn và “khi nào có thể không đúng”; cho phản hồi Đúng/Một phần/Không đúng [S8].'],
        ['Thiếu kiểm soát', 'Bất định làm tăng xu hướng nhìn thấy pattern và tìm thẩm quyền.', 'Trả lại agency: điều có thể quan sát, câu hỏi phản tư, hành động nhỏ; không fear upsell.'],
        ['Confirmation bias', 'Nhớ phần khớp, bỏ qua phần không khớp.', 'Cân bằng hỗ trợ/căng thẳng; gợi ý phản chứng; không biến mọi phản hồi thành “lá số vẫn đúng”.'],
        ['Authority bias', 'Tin danh xưng, huy hiệu, “chuyên gia” không có thật.', 'Chỉ dùng người thật, vai trò thật, phương pháp thật; tách engine, rule, AI và review.'],
        ['Default effect', 'Mặc định lưu/chia sẻ/marketing gây xâm phạm.', 'Private by default; marketing/AI training/third-party sharing opt-in riêng.'],
    ], [1750, 3200, 4410], font_size=8.55
)

add_h2('2.5 Pain-point map')
add_bullets([
    ('“Có thật sự tính theo dữ liệu của tôi?” ', 'Hiển thị input, timezone, phương pháp, engine version và evidence.'),
    ('“Tôi không biết chính xác giờ sinh.” ', 'Có lựa chọn “Không rõ giờ sinh”; không ép đoán; giải thích phần bị ảnh hưởng.'),
    ('“Tôi sợ bị phán xấu.” ', 'Không dùng đỏ = xấu; diễn giải xác suất, trung tính, có điều trong vùng kiểm soát.'),
    ('“Thuật ngữ quá khó.” ', 'Hai tầng: ngôn ngữ đời thường trước, thuật ngữ và giải thích sau.'),
    ('“Tôi không biết mình sẽ nhận gì sau khi trả tiền.” ', 'Sample thật, mục lục, độ dài, thời gian, giá cuối, mua một lần/không auto-renew và chính sách tạo lại.'),
    ('“Tôi ngại lộ ngày giờ nơi sinh.” ', 'Giải thích mục đích, thời hạn lưu, quyền xóa; share card ẩn dữ liệu mặc định.'),
])

page_break()

add_h1('03. Mood, personality & creative concept')
add_h2('3.1 Creative North Star')
add_lead('Thư viện tri thức Việt đương đại — với một bàn đọc riêng tư dành cho từng người.')
add_body('“Thư viện” tạo cảm giác có phương pháp và nguồn gốc. “Đương đại” loại bỏ hình ảnh tiệm bói online. “Bàn đọc riêng tư” bổ sung sự ấm áp, gần gũi và bảo mật mà một hệ tri thức quá nghiêm trang có thể thiếu.')

add_h2('3.2 Mood & tone')
add_table(
    ['Trục', 'Đích đến', 'Không đi quá xa thành'],
    [
        ['Tĩnh', 'Không gây áp lực; nhịp đọc chậm vừa đủ.', 'Lạnh lùng, xa cách.'],
        ['Sáng rõ', 'Phân cấp, chú giải, ngôn ngữ dễ hiểu.', 'Đơn giản hóa tới mức hời hợt.'],
        ['Có chiều sâu', 'Tôn trọng phương pháp và văn hóa.', 'Cổ trang, giáo điều, thần bí hóa.'],
        ['Ấm', 'Tôn trọng, không phán xét, gần người.', 'Sến, thân mật quá mức, meme hóa.'],
        ['Có căn cứ', 'Cho thấy tại sao và giới hạn.', 'Khoa học giả, khoe kỹ thuật.'],
    ], [1350, 4300, 3710], font_size=8.8
)

add_h2('3.3 Brand archetype')
add_bullets([
    ('60% Sage: ', 'giải thích được, có cấu trúc, ưu tiên sự thật và giới hạn.'),
    ('25% Caregiver: ', 'bảo vệ người dùng dễ tổn thương, không phán xét, không làm tăng lo âu.'),
    ('15% Explorer: ', 'mở thêm góc nhìn; không coi một hệ là chân lý duy nhất.'),
    ('Không dùng Magician/Oracle làm chủ đạo: ', 'tránh xây uy quyền siêu nhiên và sự lệ thuộc.'),
])

add_h2('3.4 Ứng xử của thương hiệu')
add_bullets([
    'Rõ trước, sâu sau.',
    'Căn cứ trước, kết luận sau.',
    'Khả năng trước, định mệnh không bao giờ.',
    'Một điểm son có chủ ý tốt hơn một màn hình đầy biểu tượng.',
    'Khi chưa đủ dữ liệu, nói “chưa đủ căn cứ” thay vì bù bằng văn phong mơ hồ.',
])

add_h2('3.5 Hero line đã chốt và cách dùng an toàn')
add_callout('Canonical hero', 'Lập lá số. Hiểu vận mệnh.', fill=CINNABAR_LIGHT, accent=CINNABAR)
add_body('Giữ nguyên vì đây là quyết định đã chốt trong repo và có category clarity. Tuy nhiên “hiểu vận mệnh” phải luôn được giải nghĩa bằng supporting copy trao quyền, không được dùng như lời hứa biết trước tương lai.')
add_bullets([
    ('Supporting line bắt buộc: ', '“Một con người. Nhiều hệ quy chiếu. Một bản luận giải dễ hiểu.”'),
    ('Subhead MVP: ', '“Xem lá số Tử Vi miễn phí và khám phá những điểm nổi bật bằng lời giải thích rõ ràng, gắn với căn cứ trên chính lá số của bạn.”'),
    ('Challenger để test, chưa thay canonical: ', '“Lập lá số. Hiểu mình, rõ lựa chọn.”'),
])

page_break()

add_h1('04. Tone of voice & messaging system')
add_h2('4.1 Voice principles')
add_table(
    ['Nguyên tắc', 'Cách thể hiện'],
    [
        ['Điềm tĩnh', 'Không kích động, không countdown, không dấu chấm than liên tục.'],
        ['Rõ ràng', 'Câu ngắn; thuật ngữ được giải thích ngay tại điểm dùng.'],
        ['Tôn trọng', 'Gọi “bạn”; không phán xét đạo đức, không đóng vai “thầy”.'],
        ['Có điều kiện', 'Dùng “có xu hướng”, “có thể biểu hiện”, “trong khung phương pháp này”.'],
        ['Trao quyền', 'Kết thúc bằng điều quan sát được, câu hỏi phản tư hoặc lựa chọn tiếp theo.'],
        ['Thành thật', 'Nói rõ AI, dữ liệu, mức độ tin cậy, giới hạn và nội dung chưa đủ căn cứ.'],
    ], [2200, 7160], first_col_bold=True, font_size=8.9
)

add_h2('4.2 Đại từ và thuật ngữ')
add_bullets([
    'Dùng “bạn” trong sản phẩm và báo cáo; dùng “Lá Số Việt” hoặc “chúng tôi” ở methodology, policy và support.',
    'Không gọi người dùng là “con”, “mệnh chủ”, “đương số” trong UI phổ thông. Nếu thuật ngữ chuyên môn bắt buộc, giải thích ngay.',
    'Phân biệt: Lá số = cấu trúc tính toán; Tóm tắt = insight miễn phí; Bản luận giải = sản phẩm trả phí; Căn cứ = yếu tố/rule được sử dụng.',
    'Tên phương pháp viết: Tử Vi, Bát Tự, Kinh Dịch, Bản đồ sao. Brand luôn là Lá Số Việt; domain là lasoviet.vn.',
])

add_h2('4.3 Lexicon')
add_table(
    ['Nên dùng', 'Không dùng'],
    [
        ['có xu hướng / có thể biểu hiện', 'chắc chắn / định sẵn / không thể tránh'],
        ['điểm đáng quan sát / căn cứ được sử dụng', 'đại họa / nghiệp / khắc chết'],
        ['điều bạn có thể chủ động', 'mua ngay để hóa giải'],
        ['mức độ tin cậy / giới hạn giờ sinh', 'chính xác 99% / AI tiên tri'],
        ['bản luận giải / thanh toán một lần', 'gói Đại Cát / VIP Thiên Mệnh'],
    ], [4680, 4680]
)

add_h2('4.4 Messaging hierarchy')
add_numbered([
    'Utility — lập lá số đúng dữ liệu và phương pháp.',
    'Clarity — nhận lời giải thích tiếng Việt dễ đọc.',
    'Proof — biết vì sao nhận định xuất hiện.',
    'Agency — chuyển góc nhìn thành điều có thể quan sát hoặc hành động.',
    'Vision — một hồ sơ sinh, nhiều hệ quy chiếu; chỉ nói mạnh khi sản phẩm đã có thật.',
])

add_h2('4.5 Homepage MVP copy system')
add_bullets([
    ('Eyebrow: ', 'Nền tảng lập và luận giải Tử Vi'),
    ('H1: ', 'Lập lá số. Hiểu vận mệnh.'),
    ('Subhead: ', 'Khám phá những điểm nổi bật trong lá số bằng lời giải thích rõ ràng, gắn với căn cứ — trước khi chọn một chủ đề luận giải chuyên sâu.'),
    ('Primary CTA: ', 'Lập lá số miễn phí'),
    ('Secondary CTA: ', 'Xem báo cáo mẫu'),
    ('Trust line: ', 'Không phán định tương lai. Không dùng nỗi sợ để bán hàng.'),
])

add_h2('4.6 Tone theo ngữ cảnh')
add_table(
    ['Ngữ cảnh', 'Giọng', 'Mẫu'],
    [
        ['Form', 'Cụ thể, hướng dẫn.', '“Giờ sinh giúp xác định một số cung và chu kỳ. Nếu không rõ, bạn vẫn có thể tiếp tục với kết quả giới hạn.”'],
        ['Insight', 'Gợi mở, có căn cứ.', '“Nhận định này dựa trên… và thường rõ hơn khi…”'],
        ['Paywall', 'Minh bạch, không gây sợ.', '“Mở bản luận giải đầy đủ: mục lục, mẫu, giá và thời gian nhận kết quả.”'],
        ['Error', 'Nêu vấn đề + cách sửa.', '“Giờ sinh cần nằm trong khoảng 00:00–23:59. Kiểm tra lại hoặc chọn Không rõ giờ sinh.”'],
        ['Safety', 'Điềm tĩnh, giới hạn rõ.', '“Nội dung này không thay thế tư vấn y tế, pháp lý hoặc tài chính.”'],
        ['Support', 'Chịu trách nhiệm, không đổ lỗi.', '“Chúng tôi đã ghi nhận. Bạn có thể sửa dữ liệu và tạo lại theo chính sách này…”'],
    ], [1300, 1800, 6260], font_size=8.45
)

page_break()

add_h1('05. Visual identity system')
add_h2('5.1 Visual idea')
add_lead('Giấy — Mực — Son. Một ấn bản tri thức được biên tập kỹ, không phải sân khấu bói toán.')
add_body('Khoảng trắng và typography là nhận diện chính. Navy/mực dựng cấu trúc. Cinnabar/son hoạt động như dấu triện: ít nhưng có lực. Texture và chi tiết văn hóa chỉ là lớp phụ, không được cạnh tranh với nội dung.')

add_h2('5.2 Color system')
palette_rows = [
    ['Paper 50', '#FFFDF7', 'Surface cao nhất, input, modal'],
    ['Paper 100', '#F7F1E5', 'Canvas mặc định'],
    ['Paper 200', '#EEE5D6', 'Section/callout nhẹ'],
    ['Paper 300', '#D7CDBD', 'Divider trang trí; không mang thông tin độc lập'],
    ['Control border', '#958A7C', 'Biên input/control — đạt xấp xỉ 3:1 trên Paper 100'],
    ['Ink 900', '#14263D', 'Heading, primary action'],
    ['Ink 800', '#263445', 'Body text'],
    ['Ink 600', '#5E6873', 'Secondary text — 5.04:1 trên Paper 100'],
    ['Cinnabar 700', '#A63D2F', 'Accent/CTA — 5.61:1 trên Paper 100; trắng trên màu này 6.31:1'],
    ['Link blue', '#174F7A', 'Link — 7.66:1 trên Paper 100'],
    ['Success', '#2F6F57', 'Success — 5.29:1 trên Paper 100'],
    ['Warning', '#8A5A12', 'Warning — 5.26:1 trên Paper 100'],
    ['Error', '#8F2737', 'Error; luôn kèm icon/text, không chỉ dựa vào màu'],
]
table = add_table(['Token', 'HEX', 'Vai trò'], palette_rows, [1800, 1600, 5960], font_size=8.4)
swatches = [PAPER_50, PAPER_100, PAPER_200, PAPER_300, CONTROL_BORDER, INK_900, INK_800, INK_600, CINNABAR, LINK_BLUE, GREEN, AMBER, ERROR]
for i, fill in enumerate(swatches, start=1):
    set_cell_shading(table.rows[i].cells[1], fill)
    p = table.rows[i].cells[1].paragraphs[0]
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(WHITE if fill in [INK_900, INK_800, INK_600, CINNABAR, LINK_BLUE, GREEN, AMBER, ERROR] else INK_900)

add_h3('Color rules')
add_bullets([
    'Navy/mực là màu điều hướng và primary action mặc định. Cinnabar chiếm khoảng 5–10% diện tích và tối đa một CTA accent nổi bật trong mỗi viewport/section.',
    'Không dùng cinnabar để mã hóa mọi yếu tố “xấu”; không dùng xanh/đỏ làm “tốt/xấu” duy nhất trên lá số.',
    'Paper 300 chỉ là divider trang trí. Control hoặc thông tin cần boundary dùng Control border #958A7C hoặc outline tương phản 3:1.',
    'Mọi tổ hợp production phải qua WCAG 2.2 AA; không dựa vào HEX trong guideline như thay thế test thực tế [S12].',
])

add_h2('5.3 Typography')
add_table(
    ['Vai trò', 'Typeface production', 'Fallback / dùng trong tài liệu', 'Quy tắc'],
    [
        ['Display/editorial', BRAND_DISPLAY_FONT, 'Georgia / DejaVu Serif', 'Heading, lead, quote, bài đọc dài có chủ ý; không dùng cho form và bảng dữ liệu.'],
        ['UI/body/data', BRAND_UI_FONT, 'system-ui / DejaVu Sans', 'Form, label, body, table, CTA; dùng tabular numerals cho ngày giờ và metadata.'],
    ], [1600, 2300, 2400, 3060], font_size=8.5
)
add_body('Be Vietnam Pro được thiết kế với letterform/diacritic tiếng Việt và mục tiêu đọc tốt [S13]. Source Serif 4 là typeface open-source cho nhiều cỡ, weight và ngôn ngữ; cần QA đầy đủ dấu tiếng Việt trước khi ship [S14].')

add_h3('Type scale — web')
add_table(
    ['Style', 'Desktop', 'Mobile', 'Weight'],
    [
        ['Display', '48/56', '36/44', '600'],
        ['H1', '40/48', '32/40', '600'],
        ['H2', '32/40', '28/36', '600'],
        ['H3', '24/32', '24/32', '600'],
        ['H4', '20/28', '20/28', '600'],
        ['Lead', '18/30', '18/30', '400'],
        ['Body', '16/26', '16/26', '400'],
        ['Small', '14/22', '14/22', '400/500'],
        ['Caption', '12/18', '12/18', '500'],
        ['Button', '15/20', '15/20', '600'],
    ], [2400, 2400, 2400, 2160], font_size=8.5
)
add_bullets([
    'Reading measure 60–72 ký tự; phần luận giải tối đa khoảng 720px.',
    'Không viết HOA cả câu tiếng Việt; không letter-spacing âm với heading nhiều dấu.',
    'Preload đúng weight 400/500/600; ưu tiên variable WOFF2 và subset Vietnamese.',
])

page_break()

add_h2('5.4 Grid, spacing, shape')
add_bullets([
    ('Grid: ', 'mobile 4 cột / margin 16 / gutter 16; tablet 8 cột / margin 24 / gutter 20; desktop 12 cột / margin 40 / gutter 24.'),
    ('Containers: ', 'max 1200px; reading container 720px; chart/dashboard 1200px nhưng narrative vẫn 720px.'),
    ('Spacing scale: ', '4, 8, 12, 16, 24, 32, 48, 64, 96px.'),
    ('Radius: ', '4px nhỏ; 8px control/card; 12px modal/panel; pill chỉ dùng tag/filter.'),
    ('Elevation: ', 'card thường dùng border; shadow nhẹ chỉ cho lớp nổi: 0 8px 24px rgba(20,38,61,.10).'),
    ('Touch target: ', 'target sản phẩm tối thiểu 44×44px dù WCAG 2.2 AA quy định tối thiểu 24×24px trong nhiều trường hợp [S11].'),
])

add_h2('5.5 Imagery & illustration')
add_bullets([
    ('Ưu tiên: ', 'line-art mực navy, một điểm son; tư liệu/sách/bản khắc có nguồn; ảnh đời sống đương đại tự nhiên; texture giấy 2–4%.'),
    ('Bắt buộc: ', 'asset văn hóa phải có nguồn gốc và context; icon chuyên ngành được content expert kiểm duyệt.'),
    ('Tránh: ', 'stock “người suy tư”, thầy bói, khói hương, quả cầu, vàng kim giả, AI art giả cổ, chữ Hán/Nôm không kiểm chứng, vũ trụ tím.'),
])

add_h2('5.6 Iconography')
add_bullets([
    'Bộ cơ sở: Lucide; size 16/20/24px; stroke 1.75px; round cap/join.',
    'Icon chuyên ngành vẽ riêng trên cùng grid. Không trộn outline, filled, emoji và biểu tượng trang trí trong một flow.',
    'Icon không thay label ở hành động quan trọng hoặc thuật ngữ chuyên môn.',
])

add_h2('5.7 Lá số & data visualization')
add_bullets([
    'Lá số là công cụ đọc dữ liệu, không phải poster trang trí.',
    'Desktop: toàn cảnh + chọn cung. Mobile: overview + chọn cung + panel/list; không thu nhỏ toàn bộ tới mức không đọc được.',
    'Ghi nhãn trực tiếp khi có thể; tooltip phải hoạt động bằng tap, keyboard focus và screen reader.',
    'Không mã hóa thuận lợi/căng thẳng chỉ bằng màu; dùng label, icon, pattern và giải thích ngữ cảnh.',
    'Cung cấp view bảng/list tương đương để hỗ trợ accessibility và chia sẻ.',
    'Palette series: #1B3A57, #A63D2F, #467A6B, #8A622C, #66578A, #4E7395, #875566; luôn bổ sung marker/nét/label.',
])

add_h2('5.8 Motion & dark mode')
add_bullets([
    ('Motion: ', '120ms feedback; 180ms standard; 240ms panel/modal; di chuyển 4–8px; tôn trọng prefers-reduced-motion.'),
    ('Không dùng: ', 'parallax, sao bay, lá số xoay, “mở cổng vận mệnh”, shimmer mạnh.'),
    ('Dark mode: ', 'không thuộc MVP. Chỉ làm khi có dữ liệu đọc ban đêm; không auto-invert, cần palette được thiết kế và test riêng.'),
])

page_break()

add_h1('06. Product experience principles')
add_h2('6.1 Mười nguyên tắc UX')
add_numbered([
    'Method-first acquisition — gọi đúng tên phương pháp ở H1, title và URL.',
    'Value before account — cho lập thử và thấy giá trị trước khi yêu cầu tài khoản.',
    'Small win before paywall — chart cơ bản, ba insight, một evidence explanation.',
    'Progressive disclosure — ngôn ngữ đời thường trước, thuật ngữ/rule sau.',
    'Evidence at the moment of doubt — “Vì sao?” nằm cạnh claim, không chôn trong methodology.',
    'Uncertainty is visible — unknown time, confidence và limitation không bị giấu.',
    'Agency over fatalism — mỗi phần có điều quan sát/hành động trong vùng kiểm soát.',
    'Private by default — report noindex, URL không đoán được, share ẩn dữ liệu.',
    'Commercial clarity — sample, mục lục, giá cuối, thanh toán một lần, thời gian, support và regeneration.',
    'No dependency loops — không streak, không cảnh báo vận xấu, không khuyến khích xem lặp lại để giảm lo âu.'
])

add_h2('6.2 Canonical funnel')
add_callout('Flow', 'Từ khóa/nhu cầu → trang phương pháp → nhập dữ liệu sinh → lá số miễn phí → tóm tắt có căn cứ → chọn chủ đề sâu → thanh toán → báo cáo → lưu/chia sẻ riêng tư', fill=PAPER_100, accent=INK_900)

add_h2('6.3 Form lập lá số')
add_bullets([
    'Chỉ hỏi trường bắt buộc; giải thích “Vì sao cần giờ/nơi sinh?” ngay tại chỗ.',
    'Có “Không rõ giờ sinh”; không giả định âm thầm hoặc ép nhập đại.',
    'Ngày dương lịch là mặc định; lịch âm chỉ khi engine chuyển đổi đã kiểm thử.',
    'Resolve timezone và cho người dùng xác nhận; hỗ trợ lỗi cụ thể, không validation quá sớm.',
    'Cho xem lại input trước khi tạo; không yêu cầu số điện thoại trước khi có giá trị.',
    'Nếu lập cho người khác: nhắc xin phép; cho phép xử lý tạm thời không lưu.',
])

add_h2('6.4 Free experience')
add_bullets([
    'Toàn bộ sơ đồ lá số cơ bản; dữ liệu đầu vào và phương pháp.',
    'Ba insight nổi bật, một thế mạnh, một điểm căng thẳng và chủ đề nên xem sâu.',
    'Một evidence drawer mở được; preview 10–15% nội dung trả phí, không blur giả.',
    'Phản hồi Đúng / Một phần / Không đúng; không dùng phản hồi để biến hệ thống thành “luôn đúng”.',
])

add_h2('6.5 Evidence drawer')
add_bullets([
    'Label thân thiện; vị trí liên quan trên chart; quy tắc ngắn; confidence; ảnh hưởng của giờ sinh; action có thể quan sát.',
    'Không lộ prompt hoặc chi tiết nội bộ không cần thiết; không biến evidence thành jargon dump.',
    'Nếu evidence xung đột, mô tả tension thay vì chọn kết luận tuyệt đối.',
])

add_h2('6.6 Paywall & checkout')
add_bullets([
    'Bán chiều sâu, cấu trúc và tính cá nhân — không bán “cảnh báo bị khóa”.',
    'CTA cạnh: tên SKU, nội dung nhận được, sample, độ dài, thời gian tạo, giá cuối, mua một lần/không auto-renew nếu đúng.',
    'Hiển thị phương thức thanh toán thật, trạng thái pending/success/failed, support và quy trình nhập sai/tạo lại.',
    'Ưu tiên payment method quen thuộc tại Việt Nam sau khi review provider; QR có độ phổ biến cao nhưng vẫn cần UX chống chuyển nhầm và fraud [S7].',
])

add_h2('6.7 Report experience')
add_numbered([
    'Tóm tắt cá nhân — 5–7 luận điểm.',
    'Dữ liệu, phương pháp, timezone, rule set và giới hạn.',
    'Căn cứ chính.',
    'Thế mạnh và nguồn lực.',
    'Mâu thuẫn và điểm dễ mắc kẹt.',
    'Phân tích chủ đề.',
    'Chu kỳ/thời điểm bằng ngôn ngữ xác suất.',
    'Điều trong vùng kiểm soát.',
    'Câu hỏi tự phản chiếu.',
    'Tóm tắt hành động tối đa 5 mục.',
    'Giới hạn phương pháp và disclaimer.',
])

page_break()

add_h1('07. Trust, privacy & safety by design')
add_h2('7.1 Trust stack')
add_numbered([
    'Input trust — dữ liệu rõ, sửa được, giải thích vì sao cần.',
    'Calculation trust — hệ/phương pháp, timezone, lịch, engine version.',
    'Explanation trust — claim chính có “Vì sao?”.',
    'Source trust — nguồn tri thức, ngày cập nhật, tách tính toán/diễn giải.',
    'Commercial trust — giá, deliverable, không auto-renew, support/regeneration.',
    'Privacy trust — consent, download/delete, private/noindex.',
    'Safety trust — giới hạn theo ngữ cảnh và kênh hỗ trợ.',
])

add_h2('7.2 AI disclosure')
add_callout('Copy chuẩn', 'Lá Số Việt sử dụng công cụ tính toán theo từng phương pháp và AI để tổ chức, đối chiếu và diễn giải kết quả bằng tiếng Việt. Mỗi nhận định quan trọng đều gắn với dữ liệu lá số được sử dụng. Nội dung mang tính tham khảo và tự chiêm nghiệm.', fill=PAPER_100, accent=INK_900)
add_body('Disclosure phải xuất hiện ở “Cách chúng tôi luận giải”, FAQ và gần checkout/report; không giấu trong Terms. Không dùng “chuyên gia AI chính xác 99%”, avatar hoặc chữ ký chuyên gia giả.')

add_h2('7.3 Privacy baseline')
add_body('Từ 01.01.2026, Luật Bảo vệ dữ liệu cá nhân 91/2025/QH15 có hiệu lực [S9]. Ngày sinh, giờ sinh, nơi sinh, tên và ghi chú quan hệ có thể kết hợp thành hồ sơ cá nhân nhạy cảm về đời sống; cần legal review trước launch. Guideline này là yêu cầu sản phẩm, không phải tư vấn pháp lý.')
add_bullets([
    'Consent riêng theo mục đích: vận hành dịch vụ, marketing, chia sẻ bên thứ ba và huấn luyện AI không gộp.',
    'Nói rõ dữ liệu nào được gửi tới nhà cung cấp AI, thời hạn lưu, quyền truy cập/chỉnh sửa/xóa/rút consent.',
    'Birth profile và report private by default; noindex; link không tuần tự; analytics không chứa ngày/giờ/nơi sinh hoặc nội dung report.',
    'Share card ẩn tên đầy đủ, ngày/giờ/nơi sinh và mã đơn mặc định.',
    'Không bán dữ liệu hoặc tạo hồ sơ quảng cáo từ nội dung lá số.',
    'Có age gate và policy riêng nếu xử lý dữ liệu trẻ em.',
])

add_h2('7.4 Safety red lines')
add_bullets([
    'Không dự đoán chắc chắn tử vong, bệnh nặng, tai nạn, phản bội, phá sản.',
    'Không chẩn đoán tâm lý/y khoa; không đưa lời khuyên đầu tư, pháp lý hoặc điều trị.',
    'Không dùng nỗi sợ vừa tạo ra để upsell; không “mua để hóa giải”.',
    'Không khuyến khích hỏi/xem lặp lại cùng vấn đề để giảm lo âu.',
    'Không content targeting dựa trên vulnerability suy đoán từ report.',
    'Nếu có dấu hiệu tự hại hoặc crisis, dừng diễn giải và chuyển sang hướng dẫn hỗ trợ an toàn đã được chuyên gia rà soát.',
])

add_h2('7.5 Accessibility')
add_bullets([
    'Chuẩn bắt buộc: WCAG 2.2 AA; normal text 4.5:1, large text 3:1, UI/focus 3:1 [S10][S12].',
    'Keyboard đầy đủ; focus order theo đọc; focus không bị che; tooltip hoạt động bằng focus/tap.',
    'Zoom 200%, reflow 320px, heading semantic, labels luôn hiển thị, error gần nguyên nhân.',
    'Screen reader đọc được tên cung, trạng thái chọn và narrative liên quan.',
    'Không dùng màu là tín hiệu duy nhất; reduced motion; không tự phát âm thanh.',
])

page_break()

add_h1('08. Component and content recipes')
add_h2('8.1 Buttons & links')
add_bullets([
    'Primary: Ink 900 / chữ trắng; Accent: Cinnabar / chữ trắng; tối đa một accent CTA mỗi viewport/section.',
    'Secondary: nền trong + Control border; Tertiary: text/link có underline hoặc hover background.',
    'Disabled không chỉ giảm opacity; loading giữ width, đổi label khi tác vụ dài.',
    'Link trong nội dung dài luôn underline; visited state khác biệt trong thư viện kiến thức.',
])

add_h2('8.2 Inputs')
add_bullets([
    'Tối thiểu 44px; label ở trên, placeholder không thay label.',
    'Focus: border Ink 900 + ring 3px tương phản; error = border + icon + text cụ thể.',
    'Ngày/giờ/nơi sinh có ví dụ Việt Nam; date order rõ; không ép người dùng đoán timezone.',
])

add_h2('8.3 Cards, navigation, feedback')
add_bullets([
    'Card thường dùng border; không card-in-card quá hai cấp; card có nhiều action không click toàn bộ.',
    'Active state dùng màu + weight + indicator. Tab ngang tối đa khoảng 5 mục.',
    'Toast chỉ cho xác nhận không quan trọng; lỗi cần xử lý nằm gần nguyên nhân.',
    'Modal cho quyết định tập trung, không chứa bài đọc dài; skeleton đúng layout, shimmer nhẹ.',
])

add_h2('8.4 Copy recipes — Do / Don’t')
add_table(
    ['Tình huống', 'Không dùng', 'Dùng'],
    [
        ['Insight bất lợi', '“Bạn sẽ thất bại trong công việc năm nay.”', '“Giai đoạn này có thể tăng áp lực về vai trò. Hãy quan sát ba tín hiệu sau…”'],
        ['Mô tả chung', '“Bạn mạnh mẽ nhưng đôi khi nhạy cảm.”', '“Nhận định này dựa trên [cung/sao] và thường rõ hơn khi…”'],
        ['Paywall', '“Có cảnh báo quan trọng — mở khóa ngay.”', '“Xem mục lục, báo cáo mẫu, giá và phần diễn giải đầy đủ.”'],
        ['Giờ sinh thiếu', '“Hãy chọn một giờ gần đúng.”', '“Bạn có thể tiếp tục với kết quả giới hạn hoặc bổ sung giờ sinh sau.”'],
        ['Error', '“Dữ liệu không hợp lệ.”', '“Giờ sinh cần nằm trong khoảng 00:00–23:59.”'],
        ['AI', '“AI hiểu bạn hơn chính bạn.”', '“AI hỗ trợ tổ chức và diễn giải trong giới hạn căn cứ đã cấp.”'],
    ], [1500, 3500, 4360], font_size=8.25
)

add_h2('8.5 Naming sản phẩm')
add_bullets([
    'Luận giải Bản mệnh & Tiềm năng',
    'Luận giải Tình duyên & Hôn nhân',
    'Luận giải Công việc & Tài lộc',
    'Luận giải Vận trình năm {YYYY}',
    'Không dùng: Gói Đại Cát, VIP Thiên Mệnh, Mở khóa vận số, Hóa giải vận hạn.',
])

page_break()

add_h1('09. Validation roadmap & metrics')
add_h2('9.1 Research trước khi khóa v1.1')
add_bullets([
    '18–24 phỏng vấn theo JTBD: người mới, người xem nhiều, người hoài nghi, người đã trả tiền, người biết/không biết giờ sinh; đa dạng tuổi, địa bàn và trình độ số.',
    'Concept test 3 định vị; usability test mobile form ngày/giờ/nơi sinh; comprehension test disclaimer và confidence.',
    'Test “overview + detail panel” so với list-first trên mobile; test serif cho narrative dài trên Android phổ biến.',
    'Diary study 7–14 ngày để đo phản tư hữu ích so với lệ thuộc.',
    'Red-team với chuyên gia Tử Vi/rule set, behavioral scientist hoặc clinical safety reviewer, security/privacy và luật sư dữ liệu.',
])

add_h2('9.2 Hypothesis backlog ưu tiên')
add_table(
    ['ID', 'Giả thuyết', 'Metric', 'Guardrail'],
    [
        ['H1', 'Canonical hero + agency subhead tạo trust tốt hơn challenger.', 'Form start, trust score', 'Hiểu sai là dự báo chắc chắn'],
        ['H2', 'Giải thích giờ sinh giảm nhập đại.', 'Form completion, edit rate', 'Cognitive load'],
        ['H3', 'Unknown-time mode giảm abandonment.', 'Chart created', 'Hiểu đúng giới hạn'],
        ['H4', '3 insight + evidence tăng paid intent.', 'Paid topic selected', 'Refund, “quá chung chung”'],
        ['H5', 'Sample thật tốt hơn testimonial.', 'Checkout start', 'Time/scroll, refund'],
        ['H6', 'Giá + mua một lần + policy cạnh CTA tăng payment success.', 'Payment completion', 'Chargeback/support'],
        ['H7', 'Tone xác suất giữ usefulness và giảm lo âu.', 'Helpful rating', 'Anxiety/misinterpretation'],
        ['H8', 'Share card ẩn dữ liệu tăng chia sẻ an toàn.', 'Share completion', 'Tỷ lệ bật lại PII'],
    ], [650, 3900, 2150, 2660], font_size=8.0
)

add_h2('9.3 Dashboard không chỉ tối ưu conversion')
add_bullets([
    'Form completion, evidence open, paywall engagement, payment completion và revenue per valid chart.',
    'Refund/chargeback/regeneration; complaint “bị dọa/bị phán”; support về thanh toán/quyền riêng tư.',
    'Mức hiểu sai sản phẩm là dự báo chắc chắn; self-reported anxiety; nội dung bị đánh dấu gây lo âu.',
    'Yêu cầu xóa dữ liệu; share có lộ PII; tỷ lệ người dùng dưới tuổi policy.',
    'Repeat purchase 30/90 ngày nhưng không thúc bằng dependency loop.',
])

add_h2('9.4 Ngưỡng vận hành từ repo')
add_bullets([
    'Form completion <35% → sửa form/input/timezone trước marketing.',
    'Free summary → paywall <15% → xem lại value proposition/topic matching.',
    'Checkout completion <50% → kiểm tra payment trust/UX/provider.',
    'Refund/regeneration >8% → ưu tiên input validation và report quality.',
    'Các ngưỡng là giả thuyết vận hành, không phải benchmark ngành.',
])

page_break()

add_h1('10. Governance & decision checklist')
add_h2('10.1 Non-negotiables — luôn phải pass')
add_bullets([
    'Không định vị là thầy bói AI; không hứa biết trước tương lai.',
    'Không fake expert, fake testimonial, fake scarcity, fake countdown hoặc “chính xác 99%”.',
    'Không dùng fear/vulnerability để upsell.',
    'Mọi claim quan trọng trong report có evidence key; engine/rule/AI tách lớp.',
    'AI disclosure, privacy, giới hạn và support dễ tìm.',
    'Report private/noindex; analytics không chứa dữ liệu sinh/nội dung report.',
    'WCAG 2.2 AA và mobile readability là release gate.',
])

add_h2('10.2 Brand review checklist')
add_numbered([
    'Quyết định này có giúp “hiểu mình có căn cứ” không?',
    'Có tôn trọng quyết định đã khóa trong repo không?',
    'Nó đang mô tả bằng chứng hay giả thuyết? Nhãn đã rõ chưa?',
    'Ngôn ngữ có thể bị hiểu thành định mệnh, chẩn đoán hoặc lời khuyên chuyên môn không?',
    'Người dùng có thấy căn cứ, giới hạn và lựa chọn tiếp theo không?',
    'Mặc định có bảo vệ dữ liệu và quyền tự chủ không?',
    'Có dark pattern, urgency giả, social proof giả hoặc fear cue không?',
    'Màu, font, spacing, contrast và state có theo token không?',
    'Trải nghiệm có dùng được trên mobile, keyboard, screen reader và zoom 200% không?',
    'Metric thành công có guardrail trust/safety không?',
])

add_h2('10.3 Owner & change control')
add_bullets([
    ('Founder / Brand owner: ', 'duyệt thay đổi LOCKED và hero/positioning.'),
    ('Design system owner: ', 'token, component, accessibility và visual QA.'),
    ('Content owner: ', 'voice, terminology, copy lint và source quality.'),
    ('Product/Research: ', 'hypothesis backlog, user research và guardrail metrics.'),
    ('Trust & Safety / Legal / Security: ', 'privacy, sensitive claims, crisis flow, incident response.'),
    ('Decision log: ', 'mọi thay đổi lớn ghi ID, owner, ngày, bằng chứng, tác động và điều kiện rollback.'),
])
add_callout('Definition of done', 'Một quyết định chưa “xong” chỉ vì đẹp hoặc tăng conversion. Nó chỉ xong khi đúng brand, dễ hiểu, truy nguyên được, an toàn, accessible và đo được.', fill=CINNABAR_LIGHT, accent=CINNABAR_DARK)

page_break()

add_h1('Appendix A. Copy library')
add_h2('A1. Trust microcopy')
add_bullets([
    '“Dữ liệu sinh của bạn được dùng để lập lá số và tạo phần diễn giải này.”',
    '“Bạn có thể kiểm tra, sửa hoặc xóa hồ sơ bất cứ lúc nào.”',
    '“Bản luận giải này là riêng tư và không xuất hiện trên công cụ tìm kiếm.”',
    '“Giá hiển thị là giá thanh toán một lần. Không tự động gia hạn.”',
    '“Nhận định này phụ thuộc vào độ chính xác của giờ sinh.”',
])

add_h2('A2. Evidence microcopy')
add_bullets([
    '“Vì sao có nhận định này?”',
    '“Căn cứ được sử dụng”',
    '“Mức độ tin cậy: trung bình — nhạy với giờ sinh”',
    '“Khi nào nhận định này có thể không đúng?”',
    '“Bạn có thể đối chiếu điều này với trải nghiệm thực tế.”',
])

add_h2('A3. Safety microcopy')
add_bullets([
    '“Đây là góc nhìn tham khảo và tự chiêm nghiệm, không phải kết luận về con người bạn.”',
    '“Nội dung không thay thế tư vấn y tế, pháp lý, tài chính hoặc sức khỏe tâm thần.”',
    '“Nếu phần này làm bạn lo lắng, hãy tạm dừng và trao đổi với người bạn tin cậy hoặc chuyên gia phù hợp.”',
])

add_h2('A4. Payment & regeneration')
add_bullets([
    '“Bạn sẽ nhận: [mục lục], khoảng [độ dài], giao trong [thời gian], thanh toán một lần.”',
    '“Nếu nhập sai dữ liệu, bạn có thể sửa và yêu cầu tạo lại theo chính sách này.”',
    '“Thanh toán đang được xác nhận. Không cần thực hiện lại.”',
    '“Thanh toán chưa hoàn tất. Tài khoản của bạn chưa bị trừ tiền / hãy kiểm tra trạng thái tại…” — chỉ dùng khi đúng trạng thái hệ thống.'
])

page_break()

add_h1('Appendix B. Evidence & sources')
add_body('Nguồn được dùng để khóa guideline gồm source of truth trong repo, nghiên cứu hành vi, dữ liệu thị trường Việt Nam, pháp luật hiện hành và tiêu chuẩn accessibility. Các nghiên cứu không đại diện cho toàn bộ người Việt được ghi rõ là tín hiệu hoặc cơ chế, không phải chân dung văn hóa.')

sources = [
    ('S1', 'Lá Số Việt — MASTER_CONCEPT.md', 'https://github.com/harris1111/lasoviet.vn/blob/master/MASTER_CONCEPT.md'),
    ('S2', 'Brand, Positioning & Trust — repo', 'https://github.com/harris1111/lasoviet.vn/blob/master/docs/02-brand-and-positioning.md'),
    ('S3', 'Evidence & Audience Insights — repo', 'https://github.com/harris1111/lasoviet.vn/blob/master/docs/01-evidence-and-insights.md'),
    ('S4', 'CESifo Working Paper 11272 — Astrology and Matrimony in Vietnam (2024)', 'https://www.ifo.de/en/cesifo/publications/2024/working-paper/astrology-and-matrimony-social-reinforcement-religious-beliefs'),
    ('S5', 'Drivers of consumers’ trust and online purchase decisions in Vietnam (2026)', 'https://doi.org/10.21511/im.22(1).2026.06'),
    ('S6', 'Bộ Công an — phòng, chống lừa đảo trực tuyến (2025)', 'https://bocongan.gov.vn/bai-viet/tang-cuong-hop-tac-quoc-te-trong-phong-chong-toi-pham-lua-dao-truc-tuyen-1761383230'),
    ('S7', 'NAPAS 2025 Member Organization Conference — VietQR & security', 'https://en.napas.com.vn/napas-2025-member-organization-conference-184260317124736875.htm'),
    ('S8', 'Forer (1949) — The fallacy of personal validation', 'https://pubmed.ncbi.nlm.nih.gov/18110193/'),
    ('S9', 'Luật Bảo vệ dữ liệu cá nhân 91/2025/QH15 — hiệu lực 01.01.2026', 'https://vanban.chinhphu.vn/?classid=1&docid=214590&pageid=27160&typegroupid=3'),
    ('S10', 'W3C — WCAG 2.2 Understanding', 'https://www.w3.org/WAI/WCAG22/Understanding/'),
    ('S11', 'W3C — How to Meet WCAG 2.2 / Target Size', 'https://www.w3.org/WAI/WCAG22/quickref/'),
    ('S12', 'W3C — Contrast (Minimum)', 'https://www.w3.org/WAI/WCAG22/Understanding/contrast-minimum'),
    ('S13', 'Be Vietnam Pro — official Google Fonts specimen', 'https://fonts.google.com/specimen/Be+Vietnam+Pro'),
    ('S14', 'Source Serif — official Adobe open-source repository', 'https://github.com/adobe-fonts/source-serif'),
    ('S15', 'PLOS ONE — Positive fortune telling and financial risk taking (2022)', 'https://journals.plos.org/plosone/article?id=10.1371/journal.pone.0273233'),
    ('S16', 'Self-Determination Theory — autonomy, competence, relatedness', 'https://selfdeterminationtheory.org/theory/'),
    ('S17', 'Luật Thương mại điện tử 122/2025/QH15 — hiệu lực 01.07.2026', 'https://vanban.chinhphu.vn/?classid=1&docid=216503&pageid=27160&typegroupid=3'),
    ('S18', 'Bộ Công Thương — TMĐT Việt Nam 2024: bước tiến và thách thức', 'https://moit.gov.vn/khoa-hoc-va-cong-nghe/thuong-mai-dien-tu-viet-nam-nam-2024-nhung-buoc-tien-va-thach-thuc.html'),
]
for code, title, url in sources:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f'{code}  ')
    set_font(r, UI_FONT, 9.2, bold=True, color=CINNABAR)
    add_hyperlink(p, title, url)

add_h2('Research caveats')
add_bullets([
    'Keyword volume chứng minh nhu cầu tìm kiếm/công cụ, không chứng minh willingness-to-pay hoặc nhân khẩu học.',
    'Nghiên cứu e-commerce/payment chỉ cung cấp cơ chế trust gần kề; không trực tiếp đo hành vi mua luận giải tử vi.',
    'Nghiên cứu CESifo đo hệ quả xã hội của niềm tin Tử Vi trong hôn nhân; không cho phép suy rộng thành “mọi người Việt tin tử vi”.',
    'Nghiên cứu Barnum, autonomy và fortune-telling risk là bằng chứng tâm lý chung, không phải khảo sát đại diện Việt Nam.',
    'Màu/font/token phải được test trên thiết bị, nội dung và user thực; contrast trong guideline là baseline, không thay QA production.',
])

# Core properties
props = doc.core_properties
props.title = 'Lá Số Việt — Brand & Experience Guideline v1.0'
props.subject = 'Brand, behavioral UX, visual identity, tone of voice and product experience principles'
props.author = 'Lá Số Việt'
props.keywords = 'Lá Số Việt, brand guideline, UX, visual identity, tone of voice, trust, safety'

doc.save(OUT)
print(OUT)
